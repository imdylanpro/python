# Dylan Nelson
# July 25, 2024
# import_qrcode_document_creator.py

import os
import math
from docx import Document
from docx.shared import Inches

def insert_png_into_word(file_path, word_doc_path):
    """Inserts pngs into a word document."""
    # Initialize a Word document
    doc = _word_doc_creator()

    # Get a list of all PNG files in the specified directory
    png_files = [f for f in os.listdir(file_path) if f.lower().endswith(".png")]

    table = _table_creator(png_files, doc)

    # Set the column and row index to 0 so we start at the beginning of the table.
    row_index = 0
    col_index = 0

    for png_file in png_files:
      
        # Get the specific cell in the table
        cell = table.rows[row_index].cells[col_index]
        img_path = os.path.join(file_path, png_file)

        # Add the image to the cell
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(img_path, width=Inches(1.25))

        # Center the image that is placed into the cell.
        paragraph.alignment = 1

        # Increment the column index until it reaches 3 then reset
        #  then increment the row index.
        col_index += 1
        if col_index == 3:
            col_index = 0
            row_index += 1
    
    # Save the document.
    _save_word_doc(doc, word_doc_path)

def _word_doc_creator():
    """Creates a word document with specific parameters."""
    # Initialize a Word document
    doc = Document()

    # Set the top and bottom margin to 0.5 inches
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    return doc

def _table_creator(png_files, doc):
    "Creates the table where all of the pngs are stored."
    
    # 3 columns are always desired, but the number of rows is determined by the amount of pngs
    #   and the rows needs to be a multiple of three since 3 pngs can fit in a row.
    columns = 3
    rows = (len(png_files) / 3)
    rows = math.ceil(rows)

    table = doc.add_table(rows=rows, cols=columns)
    # Disable automatic resizing
    table.autofit = False
    # Center the table
    table.alignment = 1

    # Set the preferred width of the table.
    for row in table.rows:
        # Set the height for each row.
        row.height = Inches(2.0)
        for cell in row.cells:
            # Set the width for each column.
            cell.width = Inches(2.63)
    
    # Return the table that was just created.
    return table

def _save_word_doc(doc, word_doc_path):
    """Saves the word document."""
        # Save the Word document
    try:
        doc.save(word_doc_path)
        print(f"PNG files inserted into {word_doc_path} successfully.")
    except Exception:
        print(f"Error, try a different filename. ")

def directory_creator():
    """Identifies the directories that the user wants to use, then also """
    # Specify the directory containing PNG files, and the output directory and file name.
    png_directory = input("What is the filepath of your png images? ")
    output_word_doc = input("What is the filepath where you want to store the word document? ")
    output_word_doc_name = input("What do you want your file to be named? ")
    output_word_doc = output_word_doc + '/' + output_word_doc_name + '.docx'
    return png_directory, output_word_doc

if __name__ == '__main__':
    png_directory, output_word_doc = directory_creator()
    insert_png_into_word(png_directory, output_word_doc)
